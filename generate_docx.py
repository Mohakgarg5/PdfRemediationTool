"""
generate_docx.py - Convert TEAM_DOCUMENTATION.md to a formatted Word document.

Run:
    python generate_docx.py

Output:
    TEAM_DOCUMENTATION.docx
"""
import re
from pathlib import Path

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml


def setup_styles(doc):
    """Configure document styles for a professional look."""
    style = doc.styles["Normal"]
    font = style.font
    font.name = "Calibri"
    font.size = Pt(10.5)
    font.color.rgb = RGBColor(0x2D, 0x2D, 0x2D)
    style.paragraph_format.space_after = Pt(6)
    style.paragraph_format.line_spacing = 1.15

    # Heading styles
    colors = {
        "Heading 1": (RGBColor(0x1A, 0x3C, 0x6E), Pt(22), True),
        "Heading 2": (RGBColor(0x1A, 0x3C, 0x6E), Pt(16), True),
        "Heading 3": (RGBColor(0x2E, 0x5C, 0x8A), Pt(13), True),
        "Heading 4": (RGBColor(0x2E, 0x5C, 0x8A), Pt(11.5), True),
    }
    for name, (color, size, bold) in colors.items():
        h = doc.styles[name]
        h.font.name = "Calibri"
        h.font.size = size
        h.font.color.rgb = color
        h.font.bold = bold
        h.paragraph_format.space_before = Pt(18) if name == "Heading 1" else Pt(14)
        h.paragraph_format.space_after = Pt(6)
        h.paragraph_format.keep_with_next = True

    # Title style
    title_style = doc.styles["Title"]
    title_style.font.name = "Calibri"
    title_style.font.size = Pt(28)
    title_style.font.color.rgb = RGBColor(0x1A, 0x3C, 0x6E)
    title_style.font.bold = True
    title_style.paragraph_format.space_after = Pt(4)

    # Subtitle style
    sub = doc.styles["Subtitle"]
    sub.font.name = "Calibri"
    sub.font.size = Pt(13)
    sub.font.color.rgb = RGBColor(0x5A, 0x5A, 0x5A)
    sub.font.bold = False
    sub.font.italic = True
    sub.paragraph_format.space_after = Pt(24)

    # Code style (character)
    if "Code Char" not in [s.name for s in doc.styles]:
        code_char = doc.styles.add_style("Code Char", WD_STYLE_TYPE.CHARACTER)
        code_char.font.name = "Consolas"
        code_char.font.size = Pt(9.5)
        code_char.font.color.rgb = RGBColor(0xC0, 0x39, 0x2B)

    # Code block style (paragraph)
    if "Code Block" not in [s.name for s in doc.styles]:
        code_block = doc.styles.add_style("Code Block", WD_STYLE_TYPE.PARAGRAPH)
        code_block.font.name = "Consolas"
        code_block.font.size = Pt(9)
        code_block.font.color.rgb = RGBColor(0x2D, 0x2D, 0x2D)
        code_block.paragraph_format.space_before = Pt(4)
        code_block.paragraph_format.space_after = Pt(4)
        code_block.paragraph_format.line_spacing = 1.0
        code_block.paragraph_format.left_indent = Cm(0.8)

    # List Bullet style
    lb = doc.styles["List Bullet"]
    lb.font.name = "Calibri"
    lb.font.size = Pt(10.5)
    lb.paragraph_format.space_after = Pt(3)


def add_cover_page(doc):
    """Add a professional cover page."""
    # Spacer
    for _ in range(4):
        doc.add_paragraph("")

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("VAPT")
    run.font.size = Pt(42)
    run.font.color.rgb = RGBColor(0x1A, 0x3C, 0x6E)
    run.font.bold = True
    run.font.name = "Calibri"

    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run2 = p2.add_run("PDF Accessibility Remediation Pipeline")
    run2.font.size = Pt(18)
    run2.font.color.rgb = RGBColor(0x2E, 0x5C, 0x8A)
    run2.font.name = "Calibri"

    p3 = doc.add_paragraph()
    p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p3.paragraph_format.space_before = Pt(12)
    run3 = p3.add_run("Complete Technical Documentation")
    run3.font.size = Pt(14)
    run3.font.color.rgb = RGBColor(0x5A, 0x5A, 0x5A)
    run3.font.name = "Calibri"
    run3.font.italic = True

    # Divider line
    for _ in range(3):
        doc.add_paragraph("")

    p4 = doc.add_paragraph()
    p4.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run4 = p4.add_run("_" * 50)
    run4.font.color.rgb = RGBColor(0xCC, 0xCC, 0xCC)

    doc.add_paragraph("")

    info_lines = [
        ("For the team.", True),
        ("Read this if Mohak is unavailable and you need to", False),
        ("understand, maintain, debug, or extend this system.", False),
        ("", False),
        ("Last updated: April 2026", False),
    ]
    for text, bold in info_lines:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(text)
        run.font.size = Pt(11)
        run.font.color.rgb = RGBColor(0x5A, 0x5A, 0x5A)
        run.font.bold = bold
        run.font.name = "Calibri"

    doc.add_page_break()


def set_cell_shading(cell, color_hex):
    """Set background shading on a table cell."""
    shading = parse_xml(
        f'<w:shd {nsdecls("w")} w:fill="{color_hex}" w:val="clear"/>'
    )
    cell._tc.get_or_add_tcPr().append(shading)


def add_table(doc, header_row, data_rows):
    """Add a formatted table to the document."""
    num_cols = len(header_row)
    table = doc.add_table(rows=1, cols=num_cols)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.LEFT

    # Header row
    hdr = table.rows[0]
    for i, text in enumerate(header_row):
        cell = hdr.cells[i]
        cell.text = ""
        p = cell.paragraphs[0]
        run = p.add_run(text.strip())
        run.font.bold = True
        run.font.size = Pt(9.5)
        run.font.name = "Calibri"
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        set_cell_shading(cell, "1A3C6E")

    # Data rows
    for row_data in data_rows:
        row = table.add_row()
        for i, text in enumerate(row_data):
            if i < num_cols:
                cell = row.cells[i]
                cell.text = ""
                p = cell.paragraphs[0]
                # Handle inline code in table cells
                add_rich_text(p, text.strip())
                for run in p.runs:
                    run.font.size = Pt(9.5)
                    run.font.name = "Calibri"

    # Alternating row shading
    for idx, row in enumerate(table.rows[1:], 1):
        if idx % 2 == 0:
            for cell in row.cells:
                set_cell_shading(cell, "F2F6FA")

    doc.add_paragraph("")  # spacing after table


def add_rich_text(paragraph, text):
    """Add text to a paragraph with inline formatting (bold, italic, code)."""
    # Process inline formatting: **bold**, *italic*, `code`
    pattern = re.compile(r'(\*\*(.+?)\*\*|`([^`]+)`|\*([^*]+?)\*)')
    last_end = 0

    for match in pattern.finditer(text):
        # Add text before this match
        if match.start() > last_end:
            paragraph.add_run(text[last_end:match.start()])

        if match.group(2):  # **bold**
            run = paragraph.add_run(match.group(2))
            run.font.bold = True
        elif match.group(3):  # `code`
            run = paragraph.add_run(match.group(3))
            run.font.name = "Consolas"
            run.font.size = Pt(9.5)
            run.font.color.rgb = RGBColor(0xC0, 0x39, 0x2B)
        elif match.group(4):  # *italic*
            run = paragraph.add_run(match.group(4))
            run.font.italic = True

        last_end = match.end()

    # Add remaining text
    if last_end < len(text):
        paragraph.add_run(text[last_end:])


def parse_markdown_table(lines, start_idx):
    """Parse a markdown table starting at start_idx. Returns (header, rows, end_idx)."""
    header_line = lines[start_idx].strip()
    cols = [c.strip() for c in header_line.strip("|").split("|")]

    # Skip separator line
    sep_idx = start_idx + 1
    if sep_idx >= len(lines):
        return None, None, start_idx

    data_rows = []
    idx = sep_idx + 1
    while idx < len(lines) and lines[idx].strip().startswith("|"):
        row_text = lines[idx].strip()
        row_cols = [c.strip() for c in row_text.strip("|").split("|")]
        data_rows.append(row_cols)
        idx += 1

    return cols, data_rows, idx


def convert_md_to_docx(md_path, docx_path):
    """Convert the markdown file to a formatted Word document."""
    doc = Document()

    # Page setup
    section = doc.sections[0]
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

    setup_styles(doc)
    add_cover_page(doc)

    with open(md_path, "r", encoding="utf-8") as f:
        lines = f.readlines()

    i = 0
    in_code_block = False
    code_block_lines = []
    skip_until_toc_end = False

    # Skip the first few lines (title, subtitle, date) since we have a cover page
    # Find where content actually starts (after TOC)
    while i < len(lines):
        line = lines[i]

        # Skip the document title (already on cover page)
        if line.startswith("# VAPT"):
            i += 1
            continue

        # Skip the "For the team" subtitle line
        if line.startswith("**For the team"):
            i += 1
            continue

        # Skip "Last updated" line
        if line.startswith("Last updated"):
            i += 1
            continue

        # Skip TOC section
        if line.strip() == "## Table of Contents":
            skip_until_toc_end = True
            i += 1
            continue

        if skip_until_toc_end:
            if line.strip() == "---" or (line.startswith("## ") and "Table of Contents" not in line):
                skip_until_toc_end = False
                if line.strip() == "---":
                    i += 1
                    continue
            else:
                i += 1
                continue

        # Skip horizontal rules
        if line.strip() == "---":
            i += 1
            continue

        # Code blocks
        if line.strip().startswith("```"):
            if in_code_block:
                # End code block - add collected lines
                code_text = "\n".join(code_block_lines)
                if code_text.strip():
                    # Add a light gray box effect via paragraph formatting
                    for code_line in code_block_lines:
                        p = doc.add_paragraph(style="Code Block")
                        run = p.add_run(code_line if code_line else " ")
                        run.font.name = "Consolas"
                        run.font.size = Pt(9)
                code_block_lines = []
                in_code_block = False
            else:
                in_code_block = True
                code_block_lines = []
            i += 1
            continue

        if in_code_block:
            code_block_lines.append(line.rstrip("\n"))
            i += 1
            continue

        # Tables
        if "|" in line and i + 1 < len(lines) and re.match(r"^\s*\|[\s\-:|]+\|", lines[i + 1]):
            header, rows, end_idx = parse_markdown_table(lines, i)
            if header and rows:
                add_table(doc, header, rows)
            i = end_idx
            continue

        # Headings
        heading_match = re.match(r"^(#{1,4})\s+(.*)", line)
        if heading_match:
            level = len(heading_match.group(1))
            text = heading_match.group(2).strip()
            # Clean markdown formatting from heading text
            text = re.sub(r'\*\*(.+?)\*\*', r'\1', text)
            text = re.sub(r'`(.+?)`', r'\1', text)

            if level == 1:
                doc.add_page_break()
                doc.add_heading(text, level=1)
            elif level == 2:
                # Add page break before major sections
                doc.add_page_break()
                doc.add_heading(text, level=1)
            elif level == 3:
                doc.add_heading(text, level=2)
            elif level == 4:
                doc.add_heading(text, level=3)
            i += 1
            continue

        # Bullet points
        bullet_match = re.match(r"^\s*[-*]\s+(.*)", line)
        if bullet_match:
            text = bullet_match.group(1).strip()
            p = doc.add_paragraph(style="List Bullet")
            add_rich_text(p, text)
            i += 1
            continue

        # Numbered list
        num_match = re.match(r"^\s*(\d+)\.\s+(.*)", line)
        if num_match:
            text = num_match.group(2).strip()
            p = doc.add_paragraph(style="List Number")
            add_rich_text(p, text)
            i += 1
            continue

        # Regular paragraph
        text = line.strip()
        if text:
            p = doc.add_paragraph()
            add_rich_text(p, text)

        i += 1

    # Add footer info
    doc.add_paragraph("")
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(
        "This document covers the complete system as of April 2026. "
        "If something here contradicts what you see in the code, "
        "the code is the source of truth - update this document."
    )
    run.font.size = Pt(9)
    run.font.italic = True
    run.font.color.rgb = RGBColor(0x88, 0x88, 0x88)

    doc.save(docx_path)
    print(f"Generated: {docx_path}")


if __name__ == "__main__":
    md_path = Path(__file__).parent / "TEAM_DOCUMENTATION.md"
    docx_path = Path(__file__).parent / "TEAM_DOCUMENTATION.docx"
    convert_md_to_docx(str(md_path), str(docx_path))
